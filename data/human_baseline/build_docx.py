#!/usr/bin/env python3
"""
Gera os anexos .docx do baseline humano a partir dos .md (fonte de verdade).

Saída:
  - Guia_do_Professor.docx       <- INSTRUCOES_ANOTADOR.md
  - Questionario_Professor.docx  <- perfil_anotador.md + template_coleta.md

Renderizador Markdown -> Word minimalista, suficiente para estes documentos:
títulos (#/##/###), blockquote (>), listas (- e 1.), negrito **, itálico *,
código `..`, régua (---). Roda com: python3 build_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent

INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def add_inline(paragraph, text):
    """Adiciona runs interpretando **negrito**, *itálico* e `código`."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_rule(doc):
    """Linha horizontal fina (borda inferior de um parágrafo vazio)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BBBBBB")
    borders.append(bottom)
    pPr.append(borders)


def add_quote(doc, text):
    """Bloco de citação: itálico, recuado, com barra cinza à esquerda."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "999999")
    borders.append(left)
    pPr.append(borders)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def add_table(doc, rows):
    """rows = lista de listas de células (strings). 1ª linha = cabeçalho."""
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r, cells in enumerate(rows):
        row = table.add_row().cells
        for c, cell in enumerate(cells):
            if c >= len(row):
                continue
            row[c].paragraphs[0].text = ""
            add_inline(row[c].paragraphs[0], cell)
            if r == 0:
                for run in row[c].paragraphs[0].runs:
                    run.bold = True


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def render_md(doc, md_text):
    lines = md_text.splitlines()
    i = 0
    quote_buf = []
    prose_buf = []

    def flush_quote():
        nonlocal quote_buf
        if quote_buf:
            add_quote(doc, " ".join(quote_buf).strip())
            quote_buf = []

    def flush_prose():
        nonlocal prose_buf
        if prose_buf:
            p = doc.add_paragraph()
            add_inline(p, " ".join(prose_buf).strip())
            prose_buf = []

    def is_special(ln):
        s = ln.lstrip()
        return (
            s.startswith(("#", ">", "- ", "* "))
            or bool(re.match(r"^\d+\.\s+", s))
            or bool(re.match(r"^---+$", ln.strip()))
        )

    while i < len(lines):
        line = lines[i].rstrip()

        # tabela: bloco de linhas iniciando com "|"
        if line.lstrip().startswith("|"):
            flush_prose(); flush_quote()
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = [parse_table_row(b) for b in block
                    if not re.match(r"^\s*\|[\s:|-]+\|?\s*$", b)]
            if rows:
                add_table(doc, rows)
            continue

        if line.startswith(">"):
            flush_prose()
            quote_buf.append(line.lstrip(">").strip())
            i += 1
            continue
        flush_quote()

        if not line.strip():
            flush_prose()
            i += 1
            continue

        if not is_special(line):
            prose_buf.append(line.strip())
            i += 1
            continue
        flush_prose()

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^---+$", line.strip()):
            add_rule(doc)
        elif re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^\s*[-*]\s+", "", line))
        elif re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\s*\d+\.\s+", "", line))
        i += 1
    flush_quote()
    flush_prose()


def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    return doc


def build(out_name, md_paths, page_break_between=False):
    doc = new_doc()
    for idx, md_path in enumerate(md_paths):
        if idx > 0 and page_break_between:
            doc.add_page_break()
        render_md(doc, Path(md_path).read_text(encoding="utf-8"))
    out = HERE / out_name
    doc.save(out)
    print(f"  ok  {out.relative_to(HERE.parent.parent)}")


def main():
    print("Gerando anexos .docx do baseline humano:")
    build("Guia_do_Professor.docx", [HERE / "INSTRUCOES_ANOTADOR.md"])
    build(
        "Questionario_Professor.docx",
        [HERE / "perfil_anotador.md", HERE / "template_coleta.md"],
        page_break_between=True,
    )
    build("Codebook_kappa.docx", [HERE / "CODEBOOK_kappa.md"])


if __name__ == "__main__":
    main()
