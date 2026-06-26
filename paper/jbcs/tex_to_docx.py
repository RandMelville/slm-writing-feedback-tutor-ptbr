#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Converte paper/jbcs/main.tex (fonte de verdade, classe sbc2023) para um .docx
editavel no Word (controle de alteracoes). Cobre o subconjunto de LaTeX usado
no manuscrito: secoes, listas, tabelas, lstlisting, quotation, \\citep/\\citet,
\\ref, \\href, \\textbf/\\textit/\\texttt e escapes/math comuns.

Uso: python3 paper/jbcs/tex_to_docx.py
Saida: paper/jbcs/benchmark_slm_jbcs.docx
"""
import re, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches

HERE = os.path.dirname(os.path.abspath(__file__))
TEX = os.path.join(HERE, "main.tex")
OUT = os.path.join(HERE, "benchmark_slm_jbcs.docx")

# --- mapas de citacao e referencia ---------------------------------------
CITE = {
    "abdin2024phi3": ("Abdin et al.", "2024"),
    "brasil2018lgpd": ("Brasil", "2018"),
    "brasil2025eca": ("Brasil", "2025"),
    "card1991infovis": ("Card et al.", "1991"),
    "cgibr2025tic": ("CGI.br et al.", "2025"),
    "gemma2024": ("Gemma Team", "2024"),
    "grattafiori2024llama3": ("Grattafiori et al.", "2024"),
    "kasneci2023chatgpt": ("Kasneci et al.", "2023"),
    "koch2018coesao": ("Koch", "2018"),
    "koch2020introducao": ("Koch", "2020"),
    "kochelias2016escrever": ("Koch & Elias", "2016"),
    "landis1977measurement": ("Landis & Koch", "1977"),
    "liu2025socratic": ("Liu et al.", "2025"),
    "macina2025mathtutorbench": ("Macina et al.", "2025"),
    "nielsen1993usability": ("Nielsen", "1993"),
    "paulelder2007socratic": ("Paul & Elder", "2007"),
    "tlili2023devil": ("Tlili et al.", "2023"),
    "vygotsky1978mind": ("Vygotsky", "1978"),
    "wood1976scaffolding": ("Wood et al.", "1976"),
    "wu2023style": ("Wu & Aji", "2023"),
    "yang2024qwen25": ("Yang et al.", "2024"),
    "zheng2023judging": ("Zheng et al.", "2023"),
    "machado1896cartomante": ("Machado de Assis", "1896"),
    "machado1882alienista": ("Machado de Assis", "1882"),
    "sabino1960homemnu": ("Sabino", "1960"),
    "colasanti1982tecela": ("Colasanti", "1982"),
    "almeida1975borboleta": ("Almeida", "1975"),
    "limabarreto1956cemiterio": ("Lima Barreto", "1956"),
}
REF = {
    "sec:intro": "1", "sec:related": "2", "sec:method": "3", "sec:efficiency": "4",
    "sec:conformance": "5", "sec:qualitative": "6", "sec:mediation": "7",
    "sec:threats": "8", "sec:conclusion": "9",
    "sec:instrument": "3.2", "sec:protocol": "3.3", "sec:validation": "3.4",
    "sec:criteria": "3.6", "sec:falsification": "3.8", "sec:llama-failure": "5.1",
    "sec:phi3": "5.2", "sec:robustness": "5.3", "sec:fm-kappa": "7.4",
    "sec:fm-model": "7.5", "sec:fm-role": "7.6", "app:stats": "A", "app:impl": "B",
    "tab:scenarios": "1", "tab:efficiency": "2", "tab:conformance": "3",
    "tab:robustness": "4", "tab:metalinguistic": "5", "tab:fm-freq": "6",
    "tab:fm-profiles": "7", "tab:fm-kappa": "8", "tab:fm-model": "9",
    "fig:fm": "1", "fig:fm-model": "2",
}
SUP = {"1": "¹", "2": "²", "3": "³", "***": "***", "*": "*",
       "**": "**", "a": "a"}

REFERENCES = [
    "Abdin, M., et al. (2024). Phi-3 technical report: A highly capable language model locally on your phone. arXiv:2404.14219. https://doi.org/10.48550/arXiv.2404.14219",
    "Almeida, L. M. de. (1975). O caso da borboleta Atiria (Serie Vaga-Lume). Atica. (Originalmente publicado em 1951 como Atiria, a borboleta, Melhoramentos)",
    "Brasil. (2018). Lei no 13.709, de 14 de agosto de 2018: Lei Geral de Protecao de Dados Pessoais (LGPD). Presidencia da Republica.",
    "Brasil. (2025). Lei no 15.211, de 22 de setembro de 2025: Estatuto Digital da Crianca e do Adolescente (ECA Digital). Presidencia da Republica.",
    "Card, S. K., Robertson, G. G., & Mackinlay, J. D. (1991). The information visualizer, an information workspace. In Proceedings of the SIGCHI Conference on Human Factors in Computing Systems (CHI '91) (pp. 181-186). ACM Press. https://doi.org/10.1145/108844.108874",
    "CGI.br, NIC.br, & Cetic.br. (2025). Pesquisa sobre o uso das tecnologias de informacao e comunicacao nas escolas brasileiras: TIC Educacao 2024. Comite Gestor da Internet no Brasil. https://cetic.br/pesquisa/educacao/",
    "Colasanti, M. (1982). A moca tecela. In Doze reis e a moca no labirinto do vento. Nordica.",
    "Gemma Team. (2024). Gemma 2: Improving open language models at a practical size. arXiv:2408.00118. https://doi.org/10.48550/arXiv.2408.00118",
    "Grattafiori, A., et al. (2024). The Llama 3 herd of models. arXiv:2407.21783. https://doi.org/10.48550/arXiv.2407.21783",
    "Kasneci, E., et al. (2023). ChatGPT for good? On opportunities and challenges of large language models for education. Learning and Individual Differences, 103, 102274. https://doi.org/10.1016/j.lindif.2023.102274",
    "Koch, I. G. V. (2018). A coesao textual (22nd ed.). Contexto. (Original work published 1989)",
    "Koch, I. G. V. (2020). Introducao a linguistica textual: trajetoria e grandes temas (2nd ed.). Contexto.",
    "Koch, I. G. V., & Elias, V. M. (2016). Escrever e argumentar. Contexto.",
    "Landis, J. R., & Koch, G. G. (1977). The measurement of observer agreement for categorical data. Biometrics, 33(1), 159-174. https://doi.org/10.2307/2529310",
    "Lima Barreto, A. H. de. (1956). O cemiterio dos vivos (F. de A. Barbosa, Ed.). In Obras de Lima Barreto. Brasiliense. (1a ed. em livro: Merito, 1953)",
    "Liu, Y., et al. (2025). Discerning minds or generic tutors? Evaluating instructional guidance capabilities in Socratic LLMs. arXiv:2508.06583. https://doi.org/10.48550/arXiv.2508.06583",
    "Machado de Assis, J. M. (1882). O alienista. In Papeis Avulsos. Lombaerts & C.",
    "Machado de Assis, J. M. (1896). A cartomante. In Varias Historias. Laemmert & C. (Publicado originalmente em 1884, Gazeta de Noticias)",
    "Macina, J., et al. (2025). MathTutorBench: A benchmark for measuring open-ended pedagogical capabilities of LLM tutors. In Proceedings of EMNLP 2025 (pp. 204-221). ACL. https://doi.org/10.18653/v1/2025.emnlp-main.11",
    "Nielsen, J. (1993). Usability engineering. Morgan Kaufmann.",
    "Paul, R., & Elder, L. (2007). Critical thinking: The art of Socratic questioning. Journal of Developmental Education, 31(1), 36-37.",
    "Sabino, F. (1960). O homem nu. Editora do Autor.",
    "Tlili, A., et al. (2023). What if the devil is my guardian angel: ChatGPT as a case study of using chatbots in education. Smart Learning Environments, 10(1), 15. https://doi.org/10.1186/s40561-023-00237-x",
    "Vygotsky, L. S. (1978). Mind in society: The development of higher psychological processes. Harvard University Press.",
    "Wood, D., Bruner, J. S., & Ross, G. (1976). The role of tutoring in problem solving. Journal of Child Psychology and Psychiatry, 17(2), 89-100. https://doi.org/10.1111/j.1469-7610.1976.tb00381.x",
    "Wu, M., & Aji, A. F. (2023). Style over substance: Evaluation biases for large language models. arXiv:2307.03025. https://doi.org/10.48550/arXiv.2307.03025",
    "Yang, A., et al. (2024). Qwen2.5 technical report. arXiv:2412.15115. https://doi.org/10.48550/arXiv.2412.15115",
    "Zheng, L., et al. (2023). Judging LLM-as-a-judge with MT-bench and Chatbot Arena. In Advances in Neural Information Processing Systems 36 (NeurIPS 2023). arXiv:2306.05685. https://doi.org/10.48550/arXiv.2306.05685",
]

DECL_TITLES = {
    "contributions": "Author Contributions",
    "interests": "Competing Interests",
    "acknowledgements": "Acknowledgements",
    "funding": "Funding",
    "materials": "Availability of Data and Materials",
}

# --- helpers de texto ----------------------------------------------------
def first_braced(s, start_idx):
    """Retorna (conteudo, indice_apos_fecha) do grupo {..} que comeca em start_idx."""
    depth = 0
    out = []
    i = start_idx
    while i < len(s):
        ch = s[i]
        if ch == '{':
            depth += 1
            if depth == 1:
                i += 1
                continue
        elif ch == '}':
            depth -= 1
            if depth == 0:
                return ''.join(out), i + 1
        out.append(ch)
        i += 1
    return ''.join(out), i

def strip_multicolumn(cell):
    """\\multicolumn{N}{spec}{content} -> content (spec may contain braces)."""
    cell = cell.strip()
    if not cell.startswith(r'\multicolumn'):
        return cell
    try:
        i = cell.index('{')
        _, i = first_braced(cell, i)            # {N}
        i = cell.index('{', i)
        _, i = first_braced(cell, i)            # {spec}
        i = cell.index('{', i)
        content, _ = first_braced(cell, i)      # {content}
        return content
    except ValueError:
        return cell

def render_cite(keys, paren):
    parts = []
    for k in [x.strip() for x in keys.split(',')]:
        a, y = CITE.get(k, (k, ''))
        parts.append(f"{a} ({y})" if not paren else f"{a}, {y}")
    if paren:
        return "(" + "; ".join(parts) + ")"
    return "; ".join(parts)

def math_sub(m):
    s = m.group(1)
    s = (s.replace(r'\times', '×').replace(r'\geq', '≥')
           .replace(r'\leq', '≤').replace(r'\sigma', 'σ')
           .replace(r'\mu', 'μ').replace(r'\alpha', 'α').replace(r'\kappa', 'κ')
           .replace(r'\approx', '≈').replace(r'\sim', '~')
           .replace(r'\chi^2', 'χ²').replace(r'\bullet', '•')
           .replace(r'\%', '%').replace(r'\,', ' '))
    s = re.sub(r'\\[a-zA-Z]+', '', s)
    return s

def preprocess(t):
    t = re.sub(r'\\href\{[^}]*\}\{([^}]*)\}', r'\1', t)
    t = re.sub(r'\\citet\{([^}]*)\}', lambda m: render_cite(m.group(1), False), t)
    t = re.sub(r'\\citep\{([^}]*)\}', lambda m: render_cite(m.group(1), True), t)
    t = re.sub(r'\\ref\{([^}]*)\}', lambda m: REF.get(m.group(1), '?'), t)
    t = re.sub(r'\\textsuperscript\{([^}]*)\}', lambda m: SUP.get(m.group(1), m.group(1)), t)
    t = re.sub(r'\$(.+?)\$', math_sub, t)
    t = t.replace(r'\noindent', '')
    t = t.replace(r'\%', '%').replace(r'\&', '&').replace(r'\_', '_')
    t = t.replace(r'\#', '#').replace(r'\,', ' ').replace(r'\ ', ' ')
    t = t.replace('``', '“').replace("''", '”')
    t = t.replace('---', '—').replace('--', '–')
    t = t.replace('~', ' ')
    return t

def strip_comments(text):
    """Remove comentarios LaTeX (% nao-escapado ate o fim da linha)."""
    return re.sub(r'(?m)(?<!\\)%.*$', '', text)

def fmt(text):
    """Converte string LaTeX em lista de runs (texto, bold, italic, mono)."""
    text = preprocess(text)
    runs, buf, stack = [], [], []
    bold = ital = mono = False
    i = 0
    def flush():
        if buf:
            runs.append((''.join(buf), bold, ital, mono))
            buf.clear()
    while i < len(text):
        if text.startswith(r'\textbf{', i):
            flush(); stack.append('b'); bold = True; i += 8; continue
        if text.startswith(r'\textit{', i):
            flush(); stack.append('i'); ital = True; i += 8; continue
        if text.startswith(r'\texttt{', i):
            flush(); stack.append('t'); mono = True; i += 8; continue
        c = text[i]
        if c == '{':
            stack.append('x'); i += 1; continue
        if c == '}':
            if stack:
                flush(); t = stack.pop()
                if t == 'b': bold = False
                elif t == 'i': ital = False
                elif t == 't': mono = False
            i += 1; continue
        if c == '\\':
            if i + 1 < len(text):
                buf.append(text[i + 1]); i += 2; continue
            i += 1; continue
        buf.append(c); i += 1
    flush()
    return [r for r in runs if r[0]]

def plain(text):
    return ''.join(r[0] for r in fmt(text))

# --- escrita no docx -----------------------------------------------------
def add_runs(par, runs, base_bold=False):
    for txt, b, i_, m in runs:
        r = par.add_run(txt)
        r.bold = b or base_bold
        r.italic = i_
        if m:
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)

def build():
    src = open(TEX, encoding='utf-8').read()
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(11)

    # titulo
    m = re.search(r'\\title(?:\[[^\]]*\])?\{(.+?)\}\s*\n', src, re.S)
    title = plain(m.group(1)) if m else "Paper"
    h = doc.add_heading(title, level=0)

    # autores
    authors = re.findall(r'\\affil\{\\textbf\{([^}]+)\}', src)
    if authors:
        p = doc.add_paragraph()
        p.add_run('; '.join(plain(a) for a in authors)).italic = True
        doc.add_paragraph('Graduate Program in Informatics in Education (PPGIE), '
                          'Federal University of Rio Grande do Sul (UFRGS), '
                          'Porto Alegre, RS, Brazil').italic = True

    # abstract
    ma = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', src, re.S)
    if ma:
        body = ma.group(1).replace(r'\textbf{Abstract.~}', '').strip()
        doc.add_heading('Abstract', level=1)
        add_runs(doc.add_paragraph(), fmt(body))
    mk = re.search(r'\\begin\{keywords\}(.*?)\\end\{keywords\}', src, re.S)
    if mk:
        p = doc.add_paragraph()
        p.add_run('Keywords: ').bold = True
        p.add_run(plain(mk.group(1).strip()))

    # corpo: de \section{Introduction} ate \section*{Declarations}
    bstart = src.index(r'\section{Introduction}')
    bend = src.index(r'\section*{Declarations}')
    process_body(doc, src[bstart:bend])

    # declaracoes
    doc.add_heading('Declarations', level=1)
    for env, label in DECL_TITLES.items():
        md = re.search(r'\\begin\{' + env + r'\}(.*?)\\end\{' + env + r'\}', src, re.S)
        if md:
            doc.add_heading(label, level=2)
            add_runs(doc.add_paragraph(), fmt(strip_comments(md.group(1)).strip()))

    # referencias
    doc.add_heading('References', level=1)
    for ref in REFERENCES:
        doc.add_paragraph(ref)

    doc.save(OUT)
    print("OK ->", OUT)

def process_body(doc, body):
    # remove font-size / alignment switches that are typeset-only (avoid leaking
    # e.g. "\footnotesize" as loose text between \begin{table*} and \begin{tabular})
    body = re.sub(r"(?m)^\s*\\(footnotesize|small|scriptsize|normalsize|"
                  r"large|Large|LARGE|centering)\s*$", "", body)
    lines = body.split('\n')
    mode = 'normal'
    tbuf, caption, cbuf, qbuf, fbuf = [], '', [], [], []
    list_mode = None
    tcount = [0]
    fcount = [0]

    def flush_figure():
        img, cap = None, ''
        for ln in fbuf:
            mi = re.search(r'\\includegraphics(?:\[[^\]]*\])?\{([^}]*)\}', ln)
            if mi:
                img = mi.group(1)
            ls = ln.strip()
            if ls.startswith(r'\caption{'):
                cap, _ = first_braced(ls, ls.index('{'))
        fcount[0] += 1
        if img:
            path = os.path.join(HERE, img)
            if os.path.exists(path):
                try:
                    doc.add_picture(path, width=Inches(5.8))
                except Exception:
                    pass
        if cap:
            p = doc.add_paragraph()
            p.add_run(f'Figure {fcount[0]}. ').bold = True
            add_runs(p, fmt(cap))

    def flush_table():
        rows = []
        for ln in tbuf:
            s = ln.strip()
            if not s or s.startswith(r'\hline') or s.startswith('%'):
                continue
            s = s.rstrip('\\').strip()
            rows.append([strip_multicolumn(c) for c in s.split('&')])
        if not rows:
            return
        tcount[0] += 1
        if caption:
            p = doc.add_paragraph()
            p.add_run(f'Table {tcount[0]}. ').bold = True
            add_runs(p, fmt(caption))
        ncol = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=ncol)
        table.style = 'Table Grid'
        for ri, row in enumerate(rows):
            for ci in range(ncol):
                cell = table.cell(ri, ci)
                cell.paragraphs[0].text = ''
                txt = row[ci] if ci < len(row) else ''
                add_runs(cell.paragraphs[0], fmt(txt), base_bold=(ri == 0))

    def flush_code():
        p = doc.add_paragraph()
        for j, cl in enumerate(cbuf):
            r = p.add_run(cl)
            r.font.name = 'Consolas'; r.font.size = Pt(9)
            if j < len(cbuf) - 1:
                r.add_break()

    for raw in lines:
        raw = re.sub(r'(?<!\\)%.*$', '', raw)
        s = raw.strip()
        if mode == 'code':
            if s.startswith(r'\end{lstlisting'):
                flush_code(); cbuf = []; mode = 'normal'
            else:
                cbuf.append(raw)
            continue
        if mode == 'table':
            if r'\end{tabular' in s:
                flush_table(); tbuf = []; caption = ''; mode = 'normal'
            else:
                tbuf.append(raw)
            continue
        if mode == 'figure':
            if r'\end{figure' in s:
                flush_figure(); fbuf = []; mode = 'normal'
            else:
                fbuf.append(raw)
            continue
        if mode == 'quote':
            if r'\end{quotation' in s:
                add_runs(doc.add_paragraph(style='Intense Quote'), fmt(' '.join(qbuf)))
                qbuf = []; mode = 'normal'
            else:
                if s:
                    qbuf.append(s)
            continue
        # normal
        if s.startswith(r'\begin{lstlisting'):
            mode = 'code'; cbuf = []; continue
        if s.startswith(r'\begin{tabular'):
            mode = 'table'; tbuf = []; continue
        if s.startswith(r'\begin{quotation'):
            mode = 'quote'; qbuf = []; continue
        if s.startswith(r'\begin{figure'):
            mode = 'figure'; fbuf = []; continue
        if s.startswith(r'\caption{'):
            caption, _ = first_braced(s, s.index('{')); continue
        if s.startswith(r'\begin{itemize'):
            list_mode = 'List Bullet'; continue
        if s.startswith(r'\begin{enumerate'):
            list_mode = 'List Number'; continue
        if s.startswith(r'\end{itemize}') or s.startswith(r'\end{enumerate}'):
            list_mode = None; continue
        if (not s or s.startswith(r'\label{') or s == r'\appendix'
                or s == r'\centering' or s.startswith(r'\hline')
                or s.startswith(r'\begin{table') or s.startswith(r'\end{table')
                or s.startswith(r'\begin{figure') or s.startswith(r'\end{figure')):
            continue
        if s.startswith(r'\section{') or s.startswith(r'\section*{'):
            c, _ = first_braced(s, s.index('{')); doc.add_heading(plain(c), level=1); continue
        if s.startswith(r'\subsection{'):
            c, _ = first_braced(s, s.index('{')); doc.add_heading(plain(c), level=2); continue
        if s.startswith(r'\subsubsection{'):
            c, _ = first_braced(s, s.index('{')); doc.add_heading(plain(c), level=3); continue
        if s.startswith(r'\paragraph{'):
            lead, end = first_braced(s, s.index('{'))
            rest = s[end:].strip()
            p = doc.add_paragraph()
            add_runs(p, fmt(lead), base_bold=True)
            if rest:
                p.add_run(' ')
                add_runs(p, fmt(rest))
            continue
        if s.startswith(r'\item'):
            content = s[len(r'\item'):].strip()
            p = doc.add_paragraph(style=list_mode or 'List Bullet')
            add_runs(p, fmt(content))
            continue
        add_runs(doc.add_paragraph(), fmt(s))

if __name__ == '__main__':
    build()
